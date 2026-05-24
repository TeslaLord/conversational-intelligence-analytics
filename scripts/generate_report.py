"""Generate the short submission report as a DOCX.

Usage:
    pip install python-docx
    python scripts/generate_report.py            # -> report.docx in project root
    python scripts/generate_report.py out.docx   # custom path
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt


# ---------- helpers ----------------------------------------------------------

def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for it in items:
        doc.add_paragraph(it, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = h
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.rows[i].cells[j].text = val
    doc.add_paragraph()  # spacer


# ---------- report -----------------------------------------------------------

def build(out_path: Path) -> None:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading("Contact-Centre Insights Bot — Submission Report", level=0)
    title.alignment = 0
    p(doc, "Repository: https://github.com/TeslaLord/conversational-intelligence-analytics")

    # ----- 1. Problem framing -----
    h1(doc, "1. Problem framing")
    p(doc,
      "Analyst-facing assistant over a synthetic customer-support conversations "
      "dataset (~64k conversations, multi-lingual, multi-channel). The system "
      "answers three classes of analyst question: numeric aggregates "
      "(e.g. resolution rate by intent), evidence retrieval (e.g. examples of "
      "escalation), and single-conversation deep dives. The design priorities "
      "were grounding (no fabricated conv_ids), reproducibility (deterministic "
      "planner, idempotent pipeline), and an honest evaluation loop.")

    # ----- 2. Dataset handling -----
    h1(doc, "2. Dataset handling")
    p(doc,
      "Input is a single CSV (customer_support_data.csv) where each row is one "
      "turn. The pipeline processes data in resumable batches; every successful "
      "batch appends conv_ids to data/processed_convs.txt, so re-running picks "
      "up the next slice without duplicating work.")

    h2(doc, "2.1 Cleaning & PII")
    bullets(doc, [
        "Text cleaning: split into sentences; drop tail sentences dominated by "
        "consonant-run gibberish (synthetic padding artefact). Always keep the "
        "first sentence. Cap at ~400 characters.",
        "PII hashing: customer_name and agent_name replaced with 16-char "
        "salted SHA-256 hashes using PII_HASH_SALT. The raw text column is "
        "dropped after cleaning — only clean_text survives downstream.",
        "Ordering: timestamps within a conversation are non-monotonic, so "
        "turn_index is the authoritative order. Duration / latency features "
        "are flagged with is_time_clean when timestamps are too broken to "
        "trust.",
    ])

    h2(doc, "2.2 Stratified sampling")
    p(doc,
      "Each batch is stratified at the conversation level by "
      "(outcome × primary_intent × language). Within each stratum we draw "
      "proportionally to its share of the remaining pool, with a floor of 1 "
      "per non-empty stratum. This keeps rare combinations (e.g. Escalated × "
      "dispute_charge × Hinglish) visible even at small batch sizes, so "
      "aggregates from early batches are comparable to later ones.")

    h2(doc, "2.3 Derived per-turn data")
    table(doc, ["Column", "Meaning"], [
        ["turn_sentiment", "Label (negative / neutral / positive) from "
         "cardiffnlp/twitter-xlm-roberta-base-sentiment (multilingual)."],
        ["turn_sentiment_score", "Confidence of the predicted label, 0–1."],
        ["turn_sentiment_signed", "−score / 0 / +score. Used to compute "
         "trajectory math (start, end, slope) at conversation level."],
        ["is_rephrase", "Customer turn whose embedding cosine-similarity to "
         "the previous customer turn is > 0.80 — i.e. they re-asked."],
        ["is_clarifying_q", "Agent turn ending in ‘?’ or containing phrases "
         "like ‘could you confirm / can you share’."],
        ["response_latency_min", "Minutes to the previous opposite-role turn. "
         "NaN if either timestamp is missing or the gap is negative."],
        ["turn_embedding", "384-d normalized vector from "
         "paraphrase-multilingual-MiniLM-L12-v2 (kept in memory; reused at "
         "indexing time so we never re-embed)."],
    ])

    h2(doc, "2.4 Derived per-conversation data")
    table(doc, ["Column", "Meaning"], [
        ["turn_count, customer_turn_count, agent_turn_count, rephrase_count",
         "Basic counts per conversation."],
        ["sentiment_start, sentiment_end, sentiment_slope",
         "First, last, and linear slope of turn_sentiment_signed."],
        ["overall_sentiment, negative_turn_ratio",
         "Modal sentiment and share of negative customer turns."],
        ["duration_minutes, avg_response_latency_min, p95_response_latency_min, is_time_clean",
         "Timing features; is_time_clean=False when timestamps were too "
         "broken to use."],
        ["task_completed, escalation_requested",
         "Outcome flags derived from the CSV outcome column."],
        ["dissatisfaction_score",
         "Composite = 0.5 · negative_turn_ratio + 0.2 · min(rephrase_count/3, 1) "
         "+ 0.3 · (1 − task_completed). Used both as a SQL ranking signal "
         "and as Chroma metadata."],
        ["summary",
         "1–3 factual sentences produced by the tagger LLM (cached on disk)."],
        ["topic_tags",
         "1–3 multi-label snake_case noun-phrases from the tagger LLM. "
         "Finer than the CSV’s single coarse primary_intent (e.g. login_issue "
         "may get [kyc_failure, biometric_glitch])."],
        ["agent_empathy_mean",
         "Float 1.0–5.0, averaged across agent turns by the tagger LLM."],
    ])

    h2(doc, "2.5 Vector metadata (what we store next to each embedding)")
    p(doc,
      "Every Chroma entry carries the per-conversation metadata used as a "
      "pre-filter at query time: industry, product, primary_intent, outcome, "
      "language, channel, overall_sentiment, agent_hash, dissatisfaction_score, "
      "month. The metadata is what lets the agent scope a semantic search to "
      "e.g. outcome=Escalated, industry=Fintech without ever touching the "
      "vector itself.")

    # ----- 3. System design -----
    h1(doc, "3. System design")
    p(doc,
      "Two clearly separated paths:")
    bullets(doc, [
        "Offline pipeline (cc_insights/pipeline.py): clean → stratified sample "
        "→ turn features → conversation metrics → LLM tagging → DuckDB upsert "
        "+ Parquet snapshot → Chroma upsert → progress log.",
        "Online agent (cc_insights/agent/*): planner LLM → deterministic "
        "orchestrator → typed tools (SQL, retrieval, conv detail) → "
        "synthesizer LLM with strict JSON output and id-grounding guards.",
    ])
    p(doc,
      "Storage: DuckDB warehouse (turns + conversations + four materialized "
      "analytic tables), Chroma vector store (three collections), SQLite for "
      "the LLM response cache and the chat session store, JSONL traces per "
      "query for audit.")

    # ----- 4. Agent flow -----
    h1(doc, "4. Agent flow")
    p(doc,
      "There is no open-ended ReAct loop. The planner picks one of six fixed "
      "intents, and the orchestrator runs a fixed tool sequence per intent. "
      "The synthesizer’s evidence is filtered against the union of conv_ids "
      "returned by the tools, so it cannot cite anything the tools did not see.")
    table(doc, ["Intent", "Tool sequence"], [
        ["aggregate",
         "sql.intent_success(filters) → retrieval.search(question, filters)"],
        ["retrieve_examples",
         "retrieval.search(question, filters)"],
        ["deep_dive_single_conv",
         "conv_detail.get(plan.conv_id)"],
        ["trajectory",
         "if no conv_id: retrieval.search(k=1) → conv_detail.get(...)"],
        ["cluster_topics",
         "sql.topic_frequency(filters) → retrieval.search(question, filters)"],
        ["agent_coaching",
         "sql.agents_needing_coaching(filters) → if rows, "
         "retrieval.search(filters + {agent_hash: worst_agent})"],
    ])
    p(doc,
      "Filter values for the planner are populated from the live DuckDB at "
      "startup (SELECT DISTINCT on the low-cardinality enum columns). Anything "
      "the planner returns that is not in the enum is dropped and logged — "
      "this single change eliminated most of the ‘no results found’ failures "
      "(e.g. the planner saying product=\"5G\" when the real value is "
      "\"5G Upgrade\").")

    # ----- 5. Retrieval -----
    h1(doc, "5. Retrieval in detail")
    h2(doc, "5.1 Three Chroma collections (chosen by intent)")
    table(doc, ["Collection", "Granularity", "Document text", "Embedding source"], [
        ["conv_summary", "1 per conversation (only if a tagger summary exists)",
         "The LLM summary string", "Freshly encoded summary"],
        ["conv_window", "4-turn sliding window, stride 3, within each conversation",
         "Concatenated index | role | text of the 4 turns",
         "Mean of the 4 turn embeddings, re-normalized (reused from feature stage)"],
        ["conv_highlight", "Individual customer-negative or rephrase turns",
         "The single turn’s clean_text",
         "That turn’s embedding (reused)"],
    ])
    p(doc,
      "All three collections share the same embedder used during feature "
      "extraction (paraphrase-multilingual-MiniLM-L12-v2, 384-d, cosine), so "
      "embeddings are never recomputed at indexing time.")

    h2(doc, "5.2 Query-time embedding + filter encoding")
    bullets(doc, [
        "The question is embedded with the same multilingual MiniLM at query "
        "time — keeps query and document vectors in the same space.",
        "Filters are converted to Chroma’s where-syntax: single key+value → "
        "{key: value}; single key + list → {key: {$in: [...]}}; multiple keys "
        "→ {$and: [...]}.",
        "Filter values are sanitized against the live enum list before they "
        "ever reach Chroma; this avoids silent zero-result searches.",
    ])

    h2(doc, "5.3 Three-attempt fallback")
    p(doc,
      "Over-constrained metadata filters are the #1 cause of empty result "
      "sets, so the retrieval tool retries before giving up:")
    bullets(doc, [
        "1. Requested collection (default: conv_summary) + filters.",
        "2. Requested collection without filters — rescues over-constrained "
        "where-clauses.",
        "3. conv_highlight without filters — rescues ‘needle-in-transcript’ "
        "questions whose answer lives in one turn, not the summary.",
    ])
    p(doc,
      "Each fallback is logged as a warning and surfaces in the JSONL trace, "
      "so a relaxed result set is never invisible.")

    h2(doc, "5.4 Output contract")
    p(doc,
      "Every hit is a RetrievedChunk(conv_id, chunk_type, text, score, "
      "metadata) where score = 1 − cosine_distance and metadata is the "
      "per-row metadata stored at indexing time. The orchestrator collects "
      "the conv_ids from every tool (SQL and retrieval) into allowed_ids; "
      "the synthesizer can cite no others, and any hallucinated id is dropped "
      "with a visible caveat (‘N hallucinated conv_id(s) removed’).")

    # ----- 6. Models & tools -----
    h1(doc, "6. Model and tool choices")
    table(doc, ["Component", "Choice", "Why"], [
        ["Embedder", "sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2 (384-d)",
         "Tiny (~120 MB), multilingual, fast on CPU; same embedder for index and query."],
        ["Sentiment", "cardiffnlp/twitter-xlm-roberta-base-sentiment",
         "Multilingual 3-class; matches the dataset’s language mix; widely benchmarked."],
        ["LLM (tag/synth/judge)", "OpenAI-compatible; default gpt-4o-mini",
         "Pluggable via LLM_BASE_URL; cheap; persistent disk cache de-duplicates."],
        ["Warehouse", "DuckDB",
         "Single-file, fast columnar; lets the SQL tool stay strictly "
         "parameterized (no free-form text-to-SQL)."],
        ["Vector store", "Chroma (persistent client, cosine)",
         "Native metadata filters; deterministic upsert IDs make the pipeline "
         "idempotent."],
        ["Planner / Synthesizer guardrails", "Pydantic schemas, strict JSON, "
         "enum-sanitization, id-grounded evidence",
         "Eliminates hallucinated conv_ids and over-constrained filters."],
    ])

    # ----- 7. Evaluation -----
    h1(doc, "7. Evaluation")
    p(doc,
      "eval/run_eval.py runs every golden question end-to-end and writes "
      "data/eval_results.json. For each question it records:")
    bullets(doc, [
        "ok_intent — planner classified the intent correctly.",
        "ok_keywords — required keywords appear in the answer body.",
        "ok_evidence — number of evidence items ≥ min_evidence.",
        "Judge rubric (1–5) from JUDGE_LLM_MODEL: faithfulness, relevance, "
        "completeness, caveats quality.",
        "Latency in milliseconds per question (end-to-end, including LLM and "
        "retrieval).",
    ])
    p(doc,
      "Representative observations from a 1000-conversation tagged subset:")
    bullets(doc, [
        "Successes: aggregate questions (top-N intents by dissatisfaction, "
        "resolution-rate-by-intent) consistently score 4–5 on faithfulness "
        "and relevance because they are anchored in the materialized SQL "
        "tables and the synthesizer cites real conv_ids from the retrieval "
        "side.",
        "Successes: deep-dive questions (‘why did C00012345 escalate?’) "
        "score high on completeness — the synthesizer narrates the turn-level "
        "sentiment trajectory and rephrase signals.",
        "Failures: vague topical phrasings (‘booking issues’, ‘5G "
        "problems’) used to return empty result sets before the planner-enum "
        "sanitization and the three-attempt retrieval fallback were added.",
        "Latency: dominated by the synthesizer LLM call; SQL is sub-50 ms, "
        "Chroma is ~100–300 ms, the LLM call is ~1–3 s on gpt-4o-mini. "
        "Persistent LLM cache makes re-runs near-instant.",
        "Consistency: deterministic planner + fixed tool sequence + id-grounded "
        "evidence makes the same question reproducible across runs (modulo "
        "LLM sampling); JSONL traces in data/traces/ make every answer "
        "auditable.",
    ])

    # ----- 8. Limitations -----
    h1(doc, "8. Limitations")
    bullets(doc, [
        "Self-preference risk in the LLM judge (same provider used for "
        "tagging, synthesis, and judging).",
        "LLM tagging is capped at TAGGING_MAX_CONVERSATIONS; remaining "
        "conversations rely only on silver labels + classifier-derived "
        "signals (no summary/topic_tags/empathy).",
        "No reranker on retrieval yet — TOP_K_FINAL is a simple top-N cut "
        "from the dense scorer.",
        "Streamlit UI is intentionally minimal (single-turn focus, session "
        "memory via SQLite); multi-turn follow-up resolution is on the "
        "roadmap.",
        "Synthetic-data artefacts: timestamps are non-monotonic and some "
        "turn tails are gibberish; both are mitigated in cleaning but make "
        "duration features only directionally meaningful.",
    ])

    # ----- 9. Next improvements -----
    h1(doc, "9. Next improvements")
    bullets(doc, [
        "Cross-encoder reranker over the top-K dense hits (e.g. bge-reranker-v2-m3).",
        "Hybrid retrieval: combine BM25 over clean_text with the dense scores.",
        "Separate, stronger judge model (e.g. gpt-4o) to reduce self-preference bias.",
        "Multi-turn agent memory — currently a single SQLite session keeps "
        "history but the planner does not resolve coreferences across turns.",
        "Stream synthesizer tokens to the UI to hide the dominant latency component.",
        "Add a lightweight evaluation dashboard that reads "
        "data/eval_results.json and tracks metrics over commits.",
    ])

    doc.save(str(out_path))
    print(f"wrote {out_path}")


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).resolve().parent.parent / "report.docx"
    build(out)
